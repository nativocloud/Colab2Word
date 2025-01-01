# Standard library imports
from typing import List, Dict, Any, Optional, Union, Literal
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime
import logging
import io

# External library imports
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

@dataclass
class DocumentTheme:
    """Controls the visual appearance of documents with consistent styling standards."""
    font_family: str = field(default='DEFAULT')
    heading1_size: int = 14
    heading2_size: int = 13
    body_size: int = 11
    primary_color: tuple = (0, 0, 0)
    paragraph_spacing: int = 6
    line_spacing: float = 1.15
    _default_fonts: List[str] = field(default_factory=lambda: [
        'Calibri',
        'Arial',
        'DejaVu Sans',
        'Helvetica',
        'sans-serif'
    ])

    def __post_init__(self):
        """Initializes font settings after class creation."""
        self.using_default_fonts = (self.font_family == 'DEFAULT')
        self.available_fonts = (self._default_fonts if self.using_default_fonts
                              else [self.font_family] + self._default_fonts)

    def get_plot_font_settings(self) -> Dict[str, Any]:
        """Creates matplotlib-compatible font configuration."""
        return {
            'font.family': 'sans-serif',
            'font.sans-serif': self.available_fonts,
            'font.size': self.body_size,
            'axes.titlesize': self.heading2_size,
            'axes.labelsize': self.body_size
        }

@dataclass
class VisualizationSettings:
    """Manages the appearance of data visualizations."""
    figure_size: tuple = (10, 6)
    dpi: int = 300
    show_grid: bool = True
    grid_alpha: float = 0.3
    plot_settings: Dict[str, Any] = field(default_factory=lambda: {
        'title_size': 14,
        'label_size': 12,
        'tick_size': 10,
        'show_values': True,
        'value_format': '.1f'
    })

class DocumentSection:
    """Represents a distinct content section within the document."""
    def __init__(
        self,
        content,
        content_type=None,
        header=None,
        description=None,
        style_options=None,
        page_break_before=False,
        orientation: Literal['portrait', 'landscape'] = 'portrait'
    ):
        """Initialize a document section with intelligent type detection."""
        # Intelligent type detection if not specified
        if content_type is None:
            if isinstance(content, str):
                content_type = 'text'
            elif isinstance(content, pd.DataFrame):
                content_type = 'table'
            elif isinstance(content, (plt.Figure, sns.axisgrid.FacetGrid)):
                content_type = 'plot'
            else:
                raise ValueError(
                    "Unable to automatically determine content type. "
                    "Please specify content_type explicitly."
                )

        # Validate content type
        valid_types = ['text', 'table', 'plot']
        if content_type not in valid_types:
            raise ValueError(f"Invalid content type. Must be one of {valid_types}")

        self.content_type = content_type
        self.data = content
        self.header = header
        self.description = description or ''
        self.style_options = style_options or {}
        self.page_break_before = page_break_before
        self.orientation = orientation

class DocumentGenerator:
    """Comprehensive system for generating professional documents."""
    def __init__(
        self,
        theme: Optional[DocumentTheme] = None,
        viz_settings: Optional[VisualizationSettings] = None,
        output_dir: Optional[Union[str, Path]] = None,
        default_subdirectory: Optional[str] = None,
        default_filename: Optional[str] = None,
        display_output: bool = True
    ):
        """Initializes the document generator with specified settings."""
        self.theme = theme or DocumentTheme()
        self.viz_settings = viz_settings or VisualizationSettings()
        self.display_output = display_output
        self._setup_logging()

        self.output_dir = (Path(output_dir) if output_dir else Path.cwd())
        if not self.output_dir.is_absolute():
            self.output_dir = Path.cwd() / self.output_dir

        self.default_subdirectory = default_subdirectory
        self.default_filename = default_filename

        self.output_dir.mkdir(parents=True, exist_ok=True)
        if self.default_subdirectory:
            (self.output_dir / self.default_subdirectory).mkdir(
                parents=True, exist_ok=True
            )

    def _setup_logging(self):
        """Configures logging system."""
        self.logger = logging.getLogger(__name__)
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)
            self.logger.setLevel(logging.INFO)

    def generate_document(
        self,
        sections: List[DocumentSection],
        output_filename: Optional[str] = None,
        subdirectory: Optional[str] = None,
        title: Optional[str] = None,
        author: Optional[str] = None,
        create_new: bool = False
    ) -> Path:
        """Generates or updates a document with the provided sections."""
        try:
            filename = output_filename or self.default_filename
            if not filename:
                raise ValueError("No filename provided and no default filename set")

            final_dir = self.output_dir
            if subdirectory or self.default_subdirectory:
                final_dir = final_dir / (subdirectory or self.default_subdirectory)
                final_dir.mkdir(parents=True, exist_ok=True)

            output_path = final_dir / filename

            if not output_path.exists() or create_new:
                self.logger.info(f"Creating new document: {output_path}")
                doc = Document()
                if title:
                    self._setup_document_properties(doc, title, author)
            else:
                self.logger.info(f"Updating existing document: {output_path}")
                doc = Document(str(output_path))

            for section in sections:
                self._add_section(doc, section)

            doc.save(str(output_path))
            self.logger.info(f"Document saved successfully: {output_path}")

            return output_path

        except Exception as e:
            self.logger.error(f"Error generating document: {str(e)}")
            raise

    def _setup_document_properties(
        self,
        doc: Document,
        title: str,
        author: Optional[str]
    ) -> None:
        """Sets up document metadata and initial formatting."""
        core_props = doc.core_properties
        core_props.title = title
        if author:
            core_props.author = author
        core_props.created = datetime.now()

        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(self.theme.heading1_size)
        title_run.font.bold = True

        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(
            f"Generated on: {datetime.now().strftime('%B %d, %Y')}"
        )

        if author:
            author_paragraph = doc.add_paragraph()
            author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_paragraph.add_run(f"Author: {author}")

        doc.add_page_break()

    def _add_section(self, doc: Document, section: DocumentSection) -> None:
        """Adds a formatted section to the document."""
        # Display content in notebook if enabled
        if self.display_output:
            self._display_section(section)

        if section.page_break_before:
            doc.add_page_break()

        if section.header:
            heading = doc.add_paragraph(section.header, style='Heading 1')

        if section.description:
            para = doc.add_paragraph(style='Normal')
            para.add_run(section.description)

        if section.content_type == 'text':
            para = doc.add_paragraph(style='Normal')
            para.add_run(section.data)
        elif section.content_type == 'table':
            include_index = section.style_options.get('include_index', True)
            self._add_table(doc, section.data, include_index=include_index)
        elif section.content_type == 'plot':
            self._add_plot(doc, section)

    def _add_table(self, doc: Document, df: pd.DataFrame, include_index: bool = True) -> None:
        """Adds a formatted table to the document."""
        if include_index:
            df_for_table = df.reset_index()
        else:
            df_for_table = df.copy()

        table = doc.add_table(rows=len(df_for_table) + 1, cols=len(df_for_table.columns))
        table.style = 'Table Grid'

        # Format headers
        for j, column in enumerate(df_for_table.columns):
            cell = table.cell(0, j)
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles['Normal']
            run = paragraph.add_run(str(column))
            run.font.bold = True
            run.font.size = Pt(self.theme.body_size)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add data
        for i, row in enumerate(df_for_table.itertuples(index=False), 1):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                paragraph = cell.paragraphs[0]
                paragraph.style = doc.styles['Normal']

                if isinstance(value, (int, np.integer)):
                    cell_text = f"{value:,}"
                elif isinstance(value, (float, np.floating)):
                    cell_text = f"${value:,.2f}" if value >= 1000 else f"{value:.2f}"
                elif pd.isna(value):
                    cell_text = "—"
                else:
                    cell_text = str(value)

                paragraph.add_run(cell_text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_plot(self, doc: Document, section: DocumentSection) -> None:
        """Adds a plot to the document."""
        try:
            fig = section.data
            buffer = io.BytesIO()
            fig.savefig(
                buffer,
                format='png',
                dpi=300,
                bbox_inches='tight',
                facecolor='white',
                edgecolor='none'
            )
            buffer.seek(0)
            
            width = section.style_options.get('width', 6)
            doc.add_picture(buffer, width=Inches(width))
            buffer.close()
        
        except Exception as e:
            self.logger.error(f"Error adding plot to document: {str(e)}")
            raise

    def _display_section(self, section: DocumentSection) -> None:
        """Displays section content in the notebook."""
        try:
            from IPython.display import display, Markdown, HTML
            
            # Display header if present
            if section.header:
                display(Markdown(f"## {section.header}"))
            
            # Display description if present
            if section.description:
                display(Markdown(f"*{section.description}*"))
            
            # Add spacing
            display(HTML("<br>"))
            
            # Display content based on type
            if section.content_type == 'text':
                display(Markdown(section.data))
            elif section.content_type == 'table':
                if hasattr(section.data, 'style'):
                    styled_df = section.data.style.set_caption('Data Table')
                    display(styled_df)
                else:
                    display(section.data)
            elif section.content_type == 'plot':
                display(section.data)
            
            # Add spacing after content
            display(HTML("<br><hr><br>"))
            
        except ImportError:
            self.logger.warning("IPython display modules not available - skipping display")
            pass